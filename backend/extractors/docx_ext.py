"""DOCX text extraction using python-docx.

Includes paragraphs, table cells, and footers (so we can pick up
'Confidential' / 'Internal' markers commonly placed in footers).
"""
from __future__ import annotations

from pathlib import Path

from . import ExtractionResult


def extract_docx(path: Path) -> ExtractionResult:
    from docx import Document

    try:
        doc = Document(str(path))
    except (KeyError, Exception) as e:
        # Some DOCX files have corrupt internal structures (e.g., missing
        # 'word/#TOC' entries). Return a clean error rather than crashing.
        return ExtractionResult(
            text=f"[Could not open DOCX: {type(e).__name__}: {e}]",
            page_count=None,
            extraction_error=f"docx_open_failed: {type(e).__name__}: {e}",
        )

    parts: list[str] = []

    # Body paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)

    # Headers & footers (often contain confidentiality markers)
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            try:
                for p in hdr.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(f"[HEADER] {t}")
            except Exception:
                pass
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            try:
                for p in ftr.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(f"[FOOTER] {t}")
            except Exception:
                pass

    # Core properties (title, author) — surface these to the LLM as a hint
    cp = doc.core_properties
    meta_lines = []
    if cp.title:
        meta_lines.append(f"Title (metadata): {cp.title}")
    if cp.author:
        meta_lines.append(f"Author (metadata): {cp.author}")
    if cp.last_modified_by:
        meta_lines.append(f"Last modified by: {cp.last_modified_by}")
    if cp.created:
        meta_lines.append(f"Created (metadata): {cp.created}")
    if cp.modified:
        meta_lines.append(f"Modified (metadata): {cp.modified}")
    if cp.subject:
        meta_lines.append(f"Subject (metadata): {cp.subject}")
    if cp.keywords:
        meta_lines.append(f"Keywords (metadata): {cp.keywords}")
    if cp.revision:
        meta_lines.append(f"Revision (metadata): {cp.revision}")

    text = "\n".join(meta_lines + parts).strip()

    # python-docx has no direct "page count" without rendering; leave None.
    return ExtractionResult(text=text, page_count=None)